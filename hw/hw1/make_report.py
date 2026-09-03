#!/usr/bin/env python
"""Assemble the HW1 report .docx from the numbers the solution actually printed.

Plain text, simple tables, embedded figures.  Each lettered part starts on a new
page so that it can be tagged separately on Gradescope.
"""

import io
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
import sys
OUT = (r"D:\MS 2024-25\Purdue\ACS - Program - Autonomy\Courses\Semesters"
       r"\Fall'26\ECE 66100\Homework\HW1\HW1 Report ShrikrishnaRajule.docx")
if len(sys.argv) > 1:
    OUT = sys.argv[1]

SOURCE = io.open(os.path.join(HERE, "hw1_ShrikrishnaRajule.py"),
                 encoding="utf-8").read()


def snippet(name):
    """Pull one function's source out of the submission file, docstring trimmed."""
    match = re.search(r"^def %s\(.*?(?=\n\ndef |\n\n#___)" % re.escape(name),
                      SOURCE, re.S | re.M)
    body = match.group(0).rstrip()
    ##  Drop the long explanatory docstring; the report explains it in prose.
    body = re.sub(r"\n[ \t]*'''.*?'''[ \t]*\n", "\n", body, count=1, flags=re.S)
    return "\n".join(line.rstrip() for line in body.splitlines() if line.strip())


document = Document()

##  US Letter, one inch margins.
for section in document.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top", "bottom", "left", "right"):
        setattr(section, "%s_margin" % side, Inches(1))

style = document.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def heading(text, level=1):
    document.add_heading(text, level=level)


def para(text, italic=False, bold=False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    return p


def blank():
    document.add_paragraph()


def page_break():
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def code(text, size=8):
    """A monospaced block, one paragraph, tight leading."""
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(text.split("\n")):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        if index < len(text.split("\n")) - 1:
            run.add_break()
    return p


def table(rows, caption=None, widths=None):
    """A simple bordered table; rows[0] is the header."""
    t = document.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = t.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.bold = (r == 0)
    if widths:
        for c, width in enumerate(widths):
            for row in t.rows:
                row.cells[c].width = Inches(width)
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
    return t


def figure(filename, caption, width=6.0):
    document.add_picture(os.path.join(HERE, filename), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True


def todo(text):
    """A slot the author must fill, marked so it cannot be missed."""
    p = document.add_paragraph()
    run = p.add_run("[ TO FILL IN: " + text + " ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p


# ----------------------------------------------------------------- title page

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ECE 661: Computer Vision — Homework 1")
run.bold = True
run.font.size = Pt(16)

for line in ["Shrikrishna Bhagirath Rajule",
             "srajule@purdue.edu",
             "Purdue University — Fall 2026",
             "September 2, 2026"]:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line).font.size = Pt(11)

blank()

# ------------------------------------------------------- AI Assistant Usage

heading("AI Assistant Usage", 1)

p = document.add_paragraph()
run = p.add_run("AI Usage: ")
run.bold = True
run.font.size = Pt(13)
run = p.add_run("____ %")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

todo("replace ____ with your honest estimate, to the nearest per cent")

para("An AI assistant was used on this homework. What it did and what I did:")

table([
    ["Done by the AI assistant", "Done by me"],
    ["Implementations of vanishing_point_estimates, perpendicular_distance, "
     "analyse_families, fit_conic, algebraic_residuals, and the four bonus "
     "functions.",
     "Recorded all 44 pixel coordinates by hand, at 200-400 % zoom."],
    ["The two helper scripts pick_points.py and wire_points.py, which are not "
     "part of the graded solution.",
     "Took the photograph used in Task 1(e), and chose which features in it to "
     "measure."],
    ["Assembly of this document, and the explanations of the conceptual "
     "questions.",
     "Diagnosed and re-recorded a first set of photograph points that were "
     "measured on the wrong surface."],
], widths=[3.2, 3.2])

blank()
para("Every function listed in the left column says so in its own docstring in "
     "hw1_ShrikrishnaRajule.py, as course policy requires.")

page_break()

# ------------------------------------------------------------------ Task 1(a)

heading("Task 1(a): Recorded points on tiles.jpg", 1)

para("Tool used: the Matplotlib interactive window, which Section 3.2 of the "
     "handout names as an acceptable alternative to GIMP. I drove it through a "
     "small script, pick_points.py, that adds a click handler so the "
     "coordinates are written to a file rather than copied by hand, and "
     "scroll-to-zoom so that every point could be placed at 200-400 % as the "
     "handout advises.")

para("Families A and B run along the two grout directions. Family C is the "
     "tile diagonals, which are not drawn in the image: for each of those "
     "three lines I recorded a corner where four tiles meet, stepped "
     "diagonally across the grid counting tiles, and recorded the corner I "
     "arrived at.")

blank()
table([
    ["Family", "Line", "Point 1 (x, y)", "Point 2 (x, y)", "Baseline (px)"],
    ["A  tile rows", "1", "(579.2, 383.5)", "(909.4, 508.0)", "353"],
    ["A  tile rows", "2", "(510.6, 571.8)", "(1119.9, 783.9)", "645"],
    ["A  tile rows", "3", "(662.1, 160.2)", "(773.8, 209.7)", "122"],
    ["B  tile columns", "1", "(172.4, 450.5)", "(418.1, 227.2)", "332"],
    ["B  tile columns", "2", "(670.1, 624.4)", "(705.2, 338.9)", "288"],
    ["B  tile columns", "3", "(1446.9, 748.8)", "(1118.3, 366.0)", "504"],
    ["C  tile diagonals", "1", "(451.6, 739.3)", "(1057.7, 645.1)", "613"],
    ["C  tile diagonals", "2", "(119.8, 303.8)", "(872.7, 381.9)", "757"],
    ["C  tile diagonals", "3", "(848.7, 302.2)", "(1009.9, 327.7)", "163"],
], caption="Table 1: the eighteen points recorded on tiles.jpg, in pixel "
           "coordinates (column, row), with the separation of each pair.",
   widths=[1.5, 0.6, 1.5, 1.5, 1.0])

blank()
figure("task1_recorded_points.png",
       "Figure 1: the recorded points drawn on tiles.jpg before anything was "
       "computed. This is the only check available on this step, and it is "
       "what confirms that the diagonals of family C were counted correctly.",
       width=5.2)

page_break()

# ------------------------------------------------------------------ Task 1(b)

heading("Task 1(b): Vanishing point estimates", 1)

heading("Prediction, written before computing anything", 2)
todo("your prediction: which family will give the least reliable vanishing "
     "point, and why. Write what you actually believed before you ran the "
     "code, not what the numbers below turned out to say.")

heading("Method", 2)
para("Each recorded pair of points gives a line as the join l = x1 x x2. Three "
     "lines give three pairs, and each pair meets in one estimate of the "
     "family's vanishing point, v = li x lj. The three estimates are "
     "dehomogenized first and averaged second: a homogeneous 3-vector is fixed "
     "only up to scale, so the mean of three 3-vectors is not a defined "
     "operation, and rescaling one of them would move the answer without "
     "moving any of the three points it stands for.")

code(snippet("vanishing_point_estimates"))

heading("Results", 2)
table([
    ["Family", "Estimate from lines (1,2)", "(1,3)", "(2,3)", "Spread (px)"],
    ["A  tile rows", "(7911.0, 3147.9)", "(4512.8, 1866.6)",
     "(5547.5, 2325.2)", "3631.7"],
    ["B  tile columns", "(756.8, -80.6)", "(744.5, -69.4)",
     "(754.0, -58.3)", "25.4"],
    ["C  tile diagonals", "(1999.3, 498.8)", "(2045.7, 491.5)",
     "(2266.5, 526.5)", "269.5"],
], caption="Table 2: the three vanishing point estimates per family and their "
           "spread, the diagonal of their bounding box.",
   widths=[1.35, 1.5, 1.35, 1.35, 0.95])

blank()
table([
    ["Family", "Mean vanishing point", "Distance from centre (800, 600)"],
    ["A  tile rows", "(5990.4, 2446.6)", "5509.1"],
    ["B  tile columns", "(751.8, -69.5)", "671.2"],
    ["C  tile diagonals", "(2103.8, 505.6)", "1307.2"],
], caption="Table 3: each family's mean vanishing point, used everywhere in "
           "parts (c) and (d), and its distance from the image centre.",
   widths=[1.7, 2.2, 2.6])

heading("Did the prediction hold?", 2)
todo("one sentence saying whether your prediction held")

para("The ordering by spread is B (25.4 px), then C (269.5 px), then A "
     "(3631.7 px), a range of a factor of 143. The striking thing is that the "
     "spread tracks the distance to the vanishing point monotonically: 671 px "
     "out gives a spread of 25 px, 1307 px out gives 270 px, and 5509 px out "
     "gives 3632 px.")

para("The mechanism is the conditioning of the meet. The three lines of a "
     "family converge at their vanishing point, so a distant vanishing point "
     "means the three lines are very nearly parallel in the image. As the "
     "angle between two lines goes to zero, the position of li x lj becomes "
     "arbitrarily sensitive to the directions of the two lines. A one-pixel "
     "error in a recorded point rotates its line through an angle of roughly "
     "1/baseline radians, and that angular error is amplified into a "
     "positional error at the intersection in proportion to the distance to "
     "it. Family A is the family whose lines are closest to parallel on the "
     "sensor, so it is the family whose vanishing point is worst determined.")

para("A second and smaller contribution is baseline length. Line A3 spans only "
     "122 px and line C3 only 163 px, the two shortest pairs in the set, and a "
     "short baseline makes the direction of the line noisier before the "
     "conditioning effect is applied to it at all.")

para("The measurement hazard that the handout warns about for family C, that a "
     "miscounted tile puts a line outside the family, did not materialise. "
     "Figure 1 shows the diagonals lying where they were meant to lie, and "
     "family C came second rather than last.")

page_break()

# ------------------------------------------------------------------ Task 1(c)

heading("Task 1(c): Collinearity of the three vanishing points", 1)

para("Ordered by spread, best determined first: B tile columns (25.4 px), then "
     "C tile diagonals (269.5 px), then A tile rows (3631.7 px). The worst "
     "determined family is A.")

para("Three points are collinear exactly when the third lies on the line "
     "through the other two, so the test is made directly: build the line "
     "through the mean vanishing points of B and C as a join, and measure the "
     "perpendicular distance from the mean vanishing point of A to it.")

code(snippet("perpendicular_distance"))

blank()
code("line through the two best mean VPs : [-575.062, 1352.06, 526234]\n"
     "COLLINEARITY RESIDUAL (perpendicular distance) = 264.943 px")

heading("Why the distance and not the determinant", 2)

para("Three homogeneous points are collinear exactly when det[v1 v2 v3] = 0, "
     "but the magnitude of that determinant measures nothing. Each vi is fixed "
     "only up to scale and the determinant is linear in each column, so "
     "rescaling one column rescales the determinant while moving no point at "
     "all. The code prints both:")

code("det[v1 v2 v3] with each v_i as [x, y, 1]^T : 389273\n"
     "the same determinant after scaling v3 by 1e6 : 3.89273e+11")

para("Same three points, same picture, a determinant six orders of magnitude "
     "larger. The determinant can answer the question \"is this exactly "
     "zero?\", and nothing else; with measured points it is never exactly "
     "zero, so on real data it answers nothing. The perpendicular distance "
     "divides by the norm of (a, b), so every scale factor cancels and what "
     "is left is a genuine length in pixels. That is a number that can be "
     "compared against the size of the image, against my own click precision, "
     "and against the same quantity computed on a different photograph, which "
     "is exactly what part (e) does.")

heading("What the number says about the claim", 2)

para("The claim under test is that the vanishing points of different families "
     "of parallel lines lying in one plane of the world are collinear. The "
     "residual of 264.94 px is about 4.8 % of the 5509 px distance from the "
     "image centre to the vanishing point being tested, and the point itself "
     "is one whose three estimates already disagree among themselves by 3632 "
     "px. On that scale the three vanishing points are collinear to within my "
     "recording error, and the claim survives the test.")

para("What would have falsified it is a residual that did not shrink as "
     "recording precision improved, or one comparable to the separation "
     "between the vanishing points themselves, which here is several thousand "
     "pixels. A residual of that size would have meant the three points were "
     "not on a common line at all, rather than being on one and measured "
     "imprecisely.")

page_break()

# ------------------------------------------------------------------ Task 1(d)

heading("Task 1(d): The vanishing line", 1)

para("The vanishing line is the join of the two best determined vanishing "
     "points, those of families B and C. The worst determined one, A, is "
     "deliberately left out of it, since it is the point whose distance to "
     "this line part (c) reports.")

code("l_v = join(VP B, VP C)\n"
     "l_v = [-575.062, 1352.06, 526234]\n"
     "l_v normalized to a^2 + b^2 = 1 : [-0.391393, 0.920224, 358.16]\n"
     "vanishing points that land inside the frame : 0 of 3")

para("All three vanishing points fall outside the frame, and the most distant "
     "is 5509 px from the image centre of a 1600 x 1200 image, so no single "
     "set of axis limits shows both the picture and that point usefully. Two "
     "panels are used.")

figure("task1_vanishing.png",
       "Figure 2: the nine recorded lines extended, the three vanishing "
       "points and the vanishing line l_v. Left, a near view holding the "
       "image and the two nearer vanishing points; right, a wide view holding "
       "all three.", width=6.4)

heading("The degenerate case", 2)

para("The case that a Cartesian implementation would have to special-case is "
     "not the vanishing point being far away. A distant vanishing point is an "
     "ordinary pair of large numbers and any implementation handles it. The "
     "case is the limit just beyond that: two lines of a family that come out "
     "exactly parallel in the image. Solving for their intersection in "
     "Cartesian coordinates means solving a 2x2 linear system whose "
     "determinant is then zero, so the Cartesian code must detect that and "
     "branch.")

para("The cross product needs no branch. l1 x l2 returns [a, b, 0], a "
     "perfectly ordinary homogeneous 3-vector that happens to have a zero "
     "third entry, and it is the ideal point in the common direction of the "
     "two lines. Nothing failed and nothing had to be detected: the point at "
     "infinity is a point like any other in this representation, which is the "
     "whole reason for adopting it.")

heading("Where a test is still needed", 2)

para("The test does not disappear; it moves. It is needed at the moment of "
     "returning to Cartesian coordinates, which is what dehom() does. There it "
     "must be decided whether the third entry is negligible, and the point "
     "worth making is that this cannot be done by asking whether x3 is small. "
     "A homogeneous vector is fixed only up to scale, so multiplying the whole "
     "vector by a thousand makes x3 a thousand times larger without moving the "
     "point it names. The test has to be relative to the other two entries, "
     "which is exactly what the supplied dehom() does.")

code(snippet("dehom"))

para("So the difficulty never lived in the intersection at all. It lives at "
     "the boundary where projective space is left and a finite pair of numbers "
     "is demanded, because that is the only place where the distinction "
     "between a finite point and an ideal point has any consequence.")

page_break()

# ------------------------------------------------------------------ Task 1(e)

heading("Task 1(e): The same measurements on my own photograph", 1)

para("The photograph is a tiled floor in a corridor, shot obliquely so that "
     "the grout lines visibly converge. It is 4624 x 2084 pixels, so the image "
     "centre is (2312, 1042). All eighteen points were recorded on the floor "
     "grout: the floor is one plane of the world, which is what the task "
     "requires, and features not lying in it, such as the recessed air vent "
     "visible on the left of the frame, were avoided.")

blank()
table([
    ["Family", "Line", "Point 1 (x, y)", "Point 2 (x, y)", "Baseline (px)"],
    ["A  first family", "1", "(2440.1, 882.2)", "(3146.4, 856.4)", "707"],
    ["A  first family", "2", "(2185.3, 1704.6)", "(3298.0, 1711.1)", "1113"],
    ["A  first family", "3", "(3114.1, 640.3)", "(3768.8, 592.0)", "656"],
    ["B  second family", "1", "(3207.7, 1182.2)", "(3085.1, 463.0)", "730"],
    ["B  second family", "2", "(3643.1, 424.3)", "(4607.4, 1714.3)", "1611"],
    ["B  second family", "3", "(2559.4, 504.9)", "(2501.4, 669.4)", "174"],
    ["C  third family", "1", "(3207.7, 1182.2)", "(3768.8, 595.2)", "812"],
    ["C  third family", "2", "(4230.0, 375.9)", "(3933.3, 808.1)", "524"],
    ["C  third family", "3", "(3117.4, 650.0)", "(3626.9, 417.8)", "560"],
], caption="Table 4: the eighteen points recorded on my own photograph.",
   widths=[1.5, 0.6, 1.5, 1.5, 1.0])

blank()
figure("task1e_recorded_points.png",
       "Figure 3: the recorded points on my photograph, checked before "
       "computing. All three families lie on the floor grout.", width=6.2)

page_break()

table([
    ["Family", "Estimate from lines (1,2)", "(1,3)", "(2,3)", "Spread (px)"],
    ["A  first family", "(-17005.0, 1592.5)", "(-2719.6, 1070.7)",
     "(-10322.0, 1631.5)", "14296.4"],
    ["B  second family", "(2911.7, -554.1)", "(2918.6, -513.8)",
     "(2926.0, -535.0)", "42.8"],
    ["C  third family", "(4871.1, -558.0)", "(4178.8, 166.3)",
     "(4462.7, 36.9)", "1001.9"],
], caption="Table 5: the three vanishing point estimates per family and their "
           "spread, on my own photograph.",
   widths=[1.35, 1.5, 1.35, 1.35, 0.95])

blank()
table([
    ["Family", "Mean vanishing point", "Distance from centre (2312, 1042)"],
    ["A  first family", "(-10015.5, 1431.6)", "12333.7"],
    ["B  second family", "(2918.8, -534.3)", "1689.1"],
    ["C  third family", "(4504.2, -118.3)", "2480.3"],
], caption="Table 6: mean vanishing points on my own photograph.",
   widths=[1.7, 2.2, 2.6])

blank()
para("Ordered by spread, best determined first: B (42.8 px), then C (1001.9 "
     "px), then A (14296.4 px). The pattern of part (b) repeats exactly: "
     "spread tracks distance to the vanishing point monotonically, 1689 px out "
     "giving 42.8 px of spread, 2480 px giving 1002 px, and 12334 px giving "
     "14296 px.")

code("l_v = join(VP B, VP C)\n"
     "l_v = [-416.024, 1585.43, 2.06137e+06]\n"
     "l_v normalized to a^2 + b^2 = 1 : [-0.253812, 0.967254, 1257.62]\n"
     "COLLINEARITY RESIDUAL (perpendicular distance) = 5184.366 px")

figure("task1e_vanishing.png",
       "Figure 4: the nine lines, three vanishing points and vanishing line "
       "on my own photograph. Again none of the three vanishing points lands "
       "inside the frame, so two panels are used.", width=6.4)

page_break()

heading("Comparison of residuals", 2)

table([
    ["Image", "Collinearity residual", "Ratio"],
    ["tiles.jpg (rendering, exact pinhole)", "264.94 px", "1x"],
    ["myphoto.jpg (real camera)", "5184.37 px", "19.6x worse"],
], caption="Table 7: the collinearity residual on the supplied rendering and "
           "on my own photograph.",
   widths=[3.0, 1.8, 1.4])

blank()
para("The residual on my photograph is 19.6 times worse, as expected. I want "
     "to be careful about attributing all of it to the lens, because on this "
     "particular photograph it is not the largest effect.")

heading("The largest effect here is the viewpoint, not the lens", 2)

para("The three lines of family A run within about 4.5 degrees of each other "
     "in the image, at directions of -2.1, +0.3 and -4.2 degrees. That family "
     "is very nearly parallel on the sensor, so by the conditioning argument "
     "of part (b) its vanishing point is badly determined: it lands at "
     "(-10015, 1432), which is 12334 px from the centre of a 4624 px wide "
     "frame, and its three estimates disagree by 14296 px. The point whose "
     "distance to the line I am reporting is that point. I photographed the "
     "floor with one tile direction almost parallel to the image plane, and "
     "that choice, rather than the optics, sets the scale of the 5184 px "
     "figure.")

para("This is also an accidental near-example of the degenerate case named in "
     "part (d). Family A is close to the limit where its vanishing point stops "
     "being a finite point and becomes an ideal point, and the enormous spread "
     "is what approaching that limit looks like in measured data.")

heading("The lens-related reason", 2)

para("tiles.jpg is a computer rendering, produced by an exact pinhole "
     "projection. Under a pinhole camera a straight line in the world images "
     "to a straight line, so the collinearity of vanishing points holds there "
     "up to my click precision and nothing else.")

para("A real camera has a lens, and a real lens has radial distortion, barrel "
     "or pincushion, growing with distance from the optical axis. Under radial "
     "distortion a straight line in the world no longer images to a straight "
     "line, it images to a slight curve. The consequences run right through "
     "the calculation: the two points I record on a grout line no longer "
     "define the true image of that line, the three lines of a family are no "
     "longer exactly concurrent, and the vanishing points move, most of all "
     "for lines recorded near the edge of the frame. Collinearity of vanishing "
     "points is a consequence of the projective camera model, and radial "
     "distortion is precisely a departure from that model, so it is the right "
     "place to look for a systematic residual.")

heading("Telling the lens apart from careless recording", 2)

para("The distinguishing property is that lens distortion is systematic and "
     "structured while careless recording is random. Any of these three tests "
     "separates them:")

for text in [
    "Record the same eighteen points a second time and recompute. Careless "
    "recording gives a different residual each time, scattered about zero. "
    "Distortion gives substantially the same residual, because the error is a "
    "fixed function of position in the frame rather than of the moment of "
    "clicking.",
    "Record only near the centre of the frame. Radial distortion falls off "
    "toward the optical axis, so the residual should shrink markedly. "
    "Recording error does not care where in the frame a point is, so it would "
    "not.",
    "Record many points along one long grout line and fit a straight line to "
    "them. Under distortion the residuals of that fit are smooth and "
    "systematically signed, all bowing the same way. Under careless recording "
    "they are random in sign.",
]:
    p = document.add_paragraph(text, style="List Number")

page_break()

# ------------------------------------------------------------------ Task 2(a)

heading("Task 2(a): Recorded points on the rim of plate.jpg", 1)

para("Eight points were recorded around the outer edge of the plate, where it "
     "meets the table, spread right around the boundary rather than bunched on "
     "one arc. Points from the inner ring, where the glaze changes shade, were "
     "excluded: that is a second and smaller circle, and points from both "
     "boundaries would not lie on any one conic.")

blank()
table([
    ["#", "x", "y", "#", "x", "y"],
    ["1", "515.4", "592.5", "5", "1111.9", "559.0"],
    ["2", "665.3", "479.3", "6", "965.2", "707.4"],
    ["3", "848.7", "418.6", "7", "778.6", "775.9"],
    ["4", "1022.6", "425.0", "8", "539.3", "760.0"],
], caption="Table 8: the eight rim points recorded on plate.jpg.",
   widths=[0.5, 1.1, 1.1, 0.5, 1.1, 1.1])

blank()
figure("task2_recorded_points.png",
       "Figure 5: the eight recorded rim points on plate.jpg.", width=5.0)

page_break()

# ------------------------------------------------------------------ Task 2(b)

heading("Task 2(b): Fitting the conic", 1)

para("Each point on a conic contributes one row [x^2, xy, y^2, x, y, 1] to a "
     "design matrix A, so that A c = 0 holds for a coefficient vector c "
     "describing a conic through all of them. Because C is fixed only up to "
     "scale there is nothing to solve for in the usual sense: what is wanted "
     "is the null direction of A, and with eight noisy measured points that is "
     "the right singular vector belonging to the smallest singular value, "
     "which is the last row of Vt from numpy.linalg.svd.")

code(snippet("fit_conic"))

blank()
code("singular values of the design matrix:\n"
     "   2.72e+06  7.109e+05  7.011e+04  349.7  53.22  0.0007365")

para("The smallest singular value is about five orders of magnitude below the "
     "next one up, so the null direction is cleanly separated and the fit is "
     "well determined.")

blank()
code("C, scaled so that ||C||_F = 1:\n"
     "   [    0.00000044     0.00000031    -0.00053853]\n"
     "   [    0.00000031     0.00000127    -0.00100968]\n"
     "   [   -0.00053853    -0.00100968     0.99999869]\n"
     "\n"
     "||C||_F = 1.000000\n"
     "max_i |x_i^T C x_i| = 4.783026e-04")

para("The residual is 4.78 x 10^-4, inside the 10^-3 to 10^-4 range the "
     "handout gives as the expected order.")

heading("What that number measures, and why it is not a distance", 2)

para("It is the algebraic residual: how far each recorded point is from "
     "satisfying the equation of the conic, not how far it lies from the curve "
     "in the image. x^T C x is a quadratic form in pixel coordinates, so its "
     "units are pixels squared times whatever units C carries, and C's scale "
     "was fixed by an arbitrary convention, ||C||_F = 1. Fixing that scale is "
     "not optional, because x^T C x is linear in C: without a stated "
     "convention the reported number could be made any value at all simply by "
     "multiplying C by a constant, which changes no geometry whatsoever.")

para("Even with the scale fixed the number is not a length. The same point at "
     "the same geometric distance from the curve gives a different x^T C x "
     "depending on where around the conic it sits, because the quadratic form "
     "grows at different rates where the curve is sharply bent than where it "
     "is flat. The genuinely geometric quantity would be the distance from the "
     "point to the nearest point of the conic, which has no closed form and is "
     "not what this fit minimises.")

page_break()

# ------------------------------------------------------------------ Task 2(c)

heading("Task 2(c): Classification", 1)

para("Rescaling the coefficient vector so that a^2 + b^2 + c^2 = 1 fixes the "
     "scale of the discriminant, which would otherwise be multiplied by the "
     "square of any rescaling.")

code("a =   0.2991140   b =   0.4134745   c =   0.8599824\n"
     "d = -727.2083607   e = -1363.4234621   f = 675172.0892562\n"
     "check a^2 + b^2 + c^2 = 1.000000000\n"
     "\n"
     "discriminant b^2 - 4ac = -0.857970\n"
     "classification: ellipse")

para("The discriminant is negative, so for a nondegenerate real conic this is "
     "an ellipse, which is what a circle seen obliquely should be. The value "
     "of -0.858 matches the -0.86 the handout gives as the expected figure for "
     "plate.jpg.")

blank()
code("the two circle constraints, neither of which survives:\n"
     "   b     =   0.4134745   (a circle needs 0)\n"
     "   a - c =  -0.5608683   (a circle needs 0)")

heading("Why a camera should not be expected to preserve those relations", 2)

para("The object is a circle in the world plane, and in a Euclidean Cartesian "
     "frame a nondegenerate conic is a circle exactly when b = 0 and a = c. "
     "That is a statement about the coefficients in one particular frame, not "
     "a property the curve carries around with it.")

para("The image plane is related to the world plane by a projective "
     "transformation H, and under it a conic transforms as C -> H^-T C H^-1. "
     "That congruence mixes all six coefficients together: every entry of the "
     "new matrix is a combination of every entry of the old one, with weights "
     "set by H. The relations b = 0 and a = c are preserved only by the "
     "similarity transformations, rotation, uniform scaling and translation, "
     "which are the ones that preserve angles. A camera in general position is "
     "not a similarity, it is a full projectivity with the plane's vanishing "
     "line at a finite distance, so there is no reason for those two "
     "particular relations among individual entries of C to survive, and they "
     "do not.")

heading("What did survive", 2)

para("Being a conic at all. x^T C x = 0 maps to an equation of the same form, "
     "so the curve is still described by a quadratic in the image.")

para("Being nondegenerate. H is invertible, so C of rank 3 maps to a matrix of "
     "rank 3.")

para("Being an ellipse specifically. The whole circle lies on the camera's "
     "side of the plane's vanishing line, so no point of it is sent to "
     "infinity and the closed bounded curve stays closed and bounded. Had the "
     "vanishing line cut through the circle, part of it would have been sent "
     "to infinity and the image would have been a hyperbola; had it touched "
     "the circle, a parabola. The classification is therefore not an accident "
     "of this photograph but a consequence of the whole plate being visible in "
     "front of the camera.")

para("Tangency survived as well, which is what part (d) checks, and so did "
     "incidence generally: a point on the curve maps to a point on the curve. "
     "What is lost is everything metric. The radius has no image, the image of "
     "the circle's centre is not the centre of the image ellipse, and the "
     "equality of all diameters is gone.")

page_break()

# ------------------------------------------------------------------ Task 2(d)

heading("Task 2(d): The tangent at a recorded point", 1)

para("The matrix C is symmetric by construction, since the xy, x and y "
     "coefficients are split evenly across the off-diagonal pairs. The tangent "
     "to the conic at a point p lying on it is l = C p, with equality up to a "
     "nonzero scale factor: l is a homogeneous line and any nonzero multiple "
     "of C p names the same line.")

code("p = (515.4, 592.5)\n"
     "l = C p = [-0.00012878032, -9.7189636e-05, 0.12419972]\n"
     "\n"
     "p^T C p = 2.414848e-04\n"
     "l^T p   = 2.414848e-04\n"
     "difference between the two = 0.000000e+00")

figure("task2_conic_and_tangent.png",
       "Figure 6: the fitted conic, the eight recorded points and the tangent "
       "at p. Left, the whole image; right, zoomed on p, where the line can be "
       "seen touching the curve without crossing it.", width=6.4)

para("The zoomed panel is the check that matters. The fitted conic follows the "
     "outer rim of the plate all the way round, and at p the tangent touches "
     "the curve and does not cross it.")

heading("Why p^T C p and l^T p must agree", 2)

para("By definition l = C p, so l^T p = (C p)^T p = p^T C^T p. C is symmetric, "
     "so C^T = C and therefore l^T p = p^T C p. The two expressions are the "
     "same expression, rewritten using the symmetry of C.")

heading("Why their agreeing verifies nothing", 2)

para("Because it is an identity. It holds for any symmetric 3 x 3 matrix and "
     "any vector p, whatever they are. It does not require that C came from a "
     "good fit, or from a fit at all; it does not require p to lie on the "
     "conic, or anywhere near the plate; it does not require the SVD to have "
     "converged on anything sensible. Substituting a random symmetric matrix "
     "and a random point reproduces the agreement exactly, to machine "
     "precision, every time.")

para("What the check does confirm is that C really is symmetric and that the "
     "matrix multiplication was coded correctly. What it cannot confirm is "
     "anything at all about whether the conic follows the rim of the plate. "
     "That is what Figure 6 is for, and in particular the zoomed panel, which "
     "is the only test in this task that could actually have failed.")

page_break()

# ---------------------------------------------------------------- Bonus (a)

heading("Bonus (a): The aiming line and the two subtended intervals", 1)

para("The aiming line for angle alpha is the join of the origin with a point "
     "one unit away in the direction alpha, which comes out as "
     "l(alpha) = [-sin alpha, cos alpha, 0]^T.")

code(snippet("aiming_line"))

blank()
code("aiming line at alpha =   0.0 deg : [ 0.000000,  1.000000,  0.000000]\n"
     "aiming line at alpha =  45.0 deg : [-0.707107,  0.707107,  0.000000]\n"
     "aiming line at alpha =  90.0 deg : [-1.000000,  0.000000,  0.000000]")

para("Alpha appears in the first two entries only. The third entry is zero "
     "because every line through the origin satisfies a*0 + b*0 + c = 0, which "
     "forces c = 0. Note also that alpha and alpha + 180 degrees give the same "
     "line up to the scale factor -1: the line carries no direction at all, "
     "which is the fact part (b) turns on.")

blank()
code("target subtends : [69.4440, 82.8750] deg\n"
     "armour subtends : [38.6598, 84.2894] deg")

para("The target's interval is contained in the armour's, but containment on "
     "its own does not establish that the target is unreachable: it shows only "
     "that any ray reaching the target also meets the armour's line somewhere, "
     "not that it meets the armour before reaching the target. The armour is "
     "met first because it is nearer. Concretely, for every angle in the "
     "target's interval the ray crosses y = 4 at an x between 0.5000 and "
     "1.5000, which lies inside the armour's span of x from 0.4 to 5.0, and "
     "y = 4 is crossed on the way to y = 8. So the shot is stopped at the "
     "armour and no direct shot at the target exists.")

blank()
code("aiming line at 45 deg met with l_inf = [0,0,1]^T :\n"
     "   [ 0.707107,  0.707107, -0.000000]  -> dehom() gives None")

para("Meeting the aiming line with the line at infinity gives an ideal point, "
     "namely [cos alpha, sin alpha, 0]^T up to scale, which is the direction "
     "of the ray itself. dehom() correctly returns None for it. Part (d) uses "
     "this. Note that [0, 0, 1]^T names two different things here: read as a "
     "point it is the origin, read as a line it is the line at infinity. That "
     "is duality, not a typographical error.")

page_break()

# ---------------------------------------------------------------- Bonus (b)

heading("Bonus (b): The naive hit test and the missing condition", 1)

code(snippet("strikes_segment"))

blank()
code("meet of the aiming line with the target's line : (2.9118, 8.0000)\n"
     "naive test (between the endpoints only)  : True\n"
     "d . q = -8.5134  ->  the meet is BEHIND the player\n"
     "corrected test (adds d . q > 0)          : False")

para("At alpha = 250 degrees, which points down and to the left, away from "
     "everything on the board, the naive test reports a hit on the target. The "
     "diagnosis is exactly the observation of part (a): alpha appears only in "
     "the first two entries of the aiming line, and l(250) and l(70) differ "
     "only by an overall factor of -1. They are the same line. A line has no "
     "direction, so it runs backwards through the target just as readily as "
     "forwards, and the meet at (2.9118, 8.0000) is a real intersection of the "
     "line with the target, just one that lies behind the player.")

para("The missing condition is a ray-direction test, not an adjustment to how "
     "the angle is bookkept. With d = (cos alpha, sin alpha)^T and q the meet "
     "in Cartesian coordinates, the shot travels forwards only when "
     "d^T q > 0. Here d^T q = -8.5134, which is negative, and the corrected "
     "test reports no hit.")

para("The inequality is strict. d^T q = 0 is the case where the meet lies on "
     "the line through the player perpendicular to the aiming direction, which "
     "for a ray starting at the player means the meet is at the player's own "
     "position, at zero range. That is not a forward hit, so the boundary case "
     "must be excluded rather than included.")

page_break()

# ---------------------------------------------------------------- Bonus (c)

heading("Bonus (c): The reflection in the wall x = 7", 1)

para("A reflection in the line l = [a, b, c]^T with a^2 + b^2 = 1 is "
     "M = I - 2 n~ l^T with n~ = [a, b, 0]^T. The reason is that ax + by + c "
     "is the signed distance to the line when a^2 + b^2 = 1, and (a, b) is the "
     "unit normal, so subtracting twice that distance along the normal is a "
     "reflection. Writing the normal as an ideal point [a, b, 0]^T is what "
     "allows the same matrix to reflect directions as well as points, which "
     "part (d) relies on.")

code(snippet("reflection_in_line"))

blank()
code("l_wall normalized to a^2 + b^2 = 1 : [-1.000000, 0.000000, 7.000000]\n"
     "\n"
     "M =                              M @ M =\n"
     "   [-1.000000  0.000000 14.000000]   [ 1.000000  0.000000  0.000000]\n"
     "   [ 0.000000  1.000000  0.000000]   [ 0.000000  1.000000  0.000000]\n"
     "   [ 0.000000  0.000000  1.000000]   [ 0.000000  0.000000  1.000000]\n"
     "\n"
     "||M @ M - I||_F = 0.00e+00")

para("M^2 = I had to be true because reflecting twice in the same line returns "
     "every point to where it started. A reflection is its own inverse, an "
     "involution, so its matrix must square to the identity.")

page_break()

# ---------------------------------------------------------------- Bonus (d)

heading("Bonus (d): The bank shot", 1)

para("The centre of the target bar is (2, 8). Reflecting it in the wall gives "
     "its mirror image, and aiming at that mirror image gives the angle whose "
     "ricochet reaches the bar.")

code("target centre                 : (2.0000, 8.0000)\n"
     "its mirror image in the wall  : (12.0000, 8.0000)\n"
     "aim at the mirror image: alpha = 33.690068 deg")

para("Tracing the ray leg by leg, rather than trusting the construction:")

blank()
table([
    ["Leg", "Quantity", "Value", "Check"],
    ["1", "crosses y = 4 at x =", "6.0000",
     "armour spans x in [0.4, 5.0]: clears it"],
    ["1", "strikes the wall at", "(7.0000, 4.6667)",
     "wall spans y in [0, 6]: hits it"],
    ["-", "incoming direction", "[0.832050, 0.554700, 0]", "an ideal point"],
    ["-", "reflected by M", "[-0.832050, 0.554700, 0]",
     "third entry still 0"],
    ["2", "lands on y = 8 at x =", "2.000000",
     "bar spans x in [1, 3]: hits it"],
    ["2", "error from bar centre", "3.11e-15", "machine precision"],
], caption="Table 9: every intermediate number of the two-leg trace.",
   widths=[0.5, 1.6, 1.7, 2.2])

blank()
para("The direction was reflected by applying M to the ideal point "
     "[cos alpha, sin alpha, 0]^T computed in part (a). The translation column "
     "of M, the 14 in the top right, multiplies the third coordinate, which is "
     "zero, so it drops out and what comes back is a pure direction rather "
     "than a translated point. A direction is an ideal point, which is why one "
     "matrix serves for both jobs, and why part (a) was worth doing.")

figure("bonus_bank_shot.png",
       "Figure 7: the two-leg bank shot on the arena. The dotted line is the "
       "aim at the mirror image of the bar centre; the dashed line is the "
       "actual path, clearing the armour, striking the wall at (7, 4.667) and "
       "returning to the centre of the target bar.", width=5.0)

page_break()

# ------------------------------------------------------------------ back matter

heading("References", 1)

for text in [
    "Course notes, ECE 661 Computer Vision, Lecture 2: homogeneous coordinates "
    "for points and lines, meet and join as cross products, ideal points and "
    "the line at infinity, conics and the tangent line. Purdue University, "
    "Fall 2026.",
    "ECE 661 Homework 1 handout, Fall 2026, including the expected check "
    "values quoted in Tasks 2(b) and 2(c).",
    "R. Hartley and A. Zisserman, Multiple View Geometry in Computer Vision, "
    "2nd edition, Cambridge University Press, 2004, chapter 2.",
    "NumPy Developers, numpy.linalg.svd documentation, "
    "https://numpy.org/doc/stable/reference/generated/numpy.linalg.svd.html",
]:
    document.add_paragraph(text, style="List Number")

document.save(OUT)
print("wrote", OUT)
print("%.1f KB" % (os.path.getsize(OUT) / 1024.0))
